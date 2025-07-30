import streamlit as st
import fitz  # PyMuPDF
import re
import tempfile
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
import requests

# ====== Config ======
TOGETHER_API_KEY = "443239d8f62ae906e1845a0d3494615e971090a4502fb5607356ccfc615a87f9"
MODEL_NAME = "mistralai/Mistral-7B-Instruct-v0.2"

# ====== Utils ======
def clean_pdf_text(text):
    text = re.sub(r'\n{2,}', '\n', text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    lines = [line for line in lines if not re.match(r'^(page|chapter)?\s*\d{1,3}$', line.lower())]
    cleaned_text = "\n".join(lines)
    cleaned_text = re.sub(r'•', '-', cleaned_text)
    cleaned_text = re.sub(r' +', ' ', cleaned_text)
    return cleaned_text

def generate_mcqs_together_ai(text, num_questions=5):
    prompt = f"""
You are a smart question paper generator bot.
Generate {num_questions} well-formatted multiple choice questions (MCQs) from the following content.
- Each question should have 4 options labeled A, B, C, D.
- Add a difficulty level tag at the end: [Easy], [Medium], or [Hard]
- Add a topic tag if identifiable, format: [Topic: <topic_name>]
- Clearly mark the correct answer using this format: "Answer: <option letter>"
- Separate each question with a blank line.

Content:
{text[:3000]}
"""
    headers = {"Authorization": f"Bearer {TOGETHER_API_KEY}"}
    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "max_tokens": 1200,
        "temperature": 0.7,
        "top_p": 0.9
    }
    try:
        response = requests.post("https://api.together.ai/v1/completions", headers=headers, json=payload)
        result = response.json()
        raw_output = result.get("choices", [{}])[0].get("text", "⚠️ No output received.")
        cleaned_output = "\n".join(line.strip() for line in raw_output.strip().splitlines())
        return cleaned_output
    except Exception as e:
        return f"❌ Error calling Together.ai API: {str(e)}"

def export_mcqs_to_pdf(mcq_text):
    mcq_blocks = [block.strip() for block in mcq_text.strip().split('\n\n') if block.strip()]
    styles = getSampleStyleSheet()
    question_style = ParagraphStyle('Question', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12, spaceAfter=6, textColor=colors.darkblue)
    option_style = ParagraphStyle('Option', parent=styles['Normal'], fontName='Helvetica', fontSize=11, leftIndent=20, spaceAfter=2, textColor=colors.black)
    answer_style = ParagraphStyle('Answer', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=11, textColor=colors.green, spaceBefore=4, spaceAfter=10)
    tag_style = ParagraphStyle('Tag', parent=styles['Normal'], fontSize=9, textColor=colors.gray, spaceAfter=10)

    elements = []
    for idx, block in enumerate(mcq_blocks, 1):
        lines = block.split('\n')
        if not lines:
            continue
        elements.append(Paragraph(f"{idx}. {lines[0]}", question_style))
        for line in lines[1:]:
            if re.match(r'^[A-D]\.', line.strip()):
                elements.append(Paragraph(line.strip(), option_style))
            elif line.strip().lower().startswith('answer:'):
                elements.append(Paragraph(line.strip(), answer_style))
            elif '[' in line and ']' in line:
                elements.append(Paragraph(line.strip(), tag_style))
        elements.append(Spacer(1, 12))
        if idx % 5 == 0:
            elements.append(PageBreak())
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    doc.build(elements)
    buffer.seek(0)
    return buffer

def export_mcqs_to_docx(mcq_text):
    doc = Document()
    doc.add_heading('Generated MCQs', 0)
    mcq_blocks = [block.strip() for block in mcq_text.strip().split('\n\n') if block.strip()]
    for idx, block in enumerate(mcq_blocks, 1):
        lines = block.split('\n')
        doc.add_paragraph(f"{idx}. {lines[0]}", style='List Number')
        for line in lines[1:]:
            if re.match(r'^[A-D]\.', line.strip()):
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.left_indent = 20
            elif line.lower().startswith('answer:'):
                p = doc.add_paragraph(line.strip())
                p.runs[0].italic = True
            elif '[' in line and ']' in line:
                doc.add_paragraph(line.strip())
        doc.add_paragraph("")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_mcqs_to_html(mcq_text):
    mcq_blocks = [block.strip() for block in mcq_text.strip().split('\n\n') if block.strip()]
    html = """
    <html><head><title>MCQ Quiz</title></head><body>
    <h1>MCQ Quiz</h1>
    <form>
    """
    for idx, block in enumerate(mcq_blocks, 1):
        lines = block.split('\n')
        html += f"<fieldset><legend><strong>Q{idx}. {lines[0]}</strong></legend>"
        for line in lines[1:]:
            if re.match(r'^[A-D]\.', line.strip()):
                option = line.strip()
                html += f"<label><input type='radio' name='q{idx}' value='{option[0]}'/> {option}</label><br>"
        tags = [l for l in lines if '[' in l and ']' in l]
        for tag in tags:
            html += f"<p style='color:gray'>{tag}</p>"
        html += "</fieldset><br>"
    html += "<input type='submit' value='Submit'></form></body></html>"
    return html

# ====== Streamlit App ======
st.title("PDF to MCQ Generator")
st.write("Upload a PDF, and generate MCQs with difficulty and topic tagging.")

uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
num_questions = st.slider("Number of MCQs", 1, 20, 5)

if uploaded_file is not None:
    # Extract text from PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_file_path = tmp_file.name
    doc = fitz.open(tmp_file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
    doc.close()
    cleaned_text = clean_pdf_text(full_text)
    st.success("PDF text extracted and cleaned.")

    if st.button("Generate MCQs"):
        with st.spinner("Generating MCQs..."):
            mcqs = generate_mcqs_together_ai(cleaned_text, num_questions=num_questions)
        st.text_area("Generated MCQs", mcqs, height=400)

        # Download buttons
        pdf_buffer = export_mcqs_to_pdf(mcqs)
        st.download_button("Download as PDF", pdf_buffer, file_name="generated_mcqs.pdf", mime="application/pdf")

        docx_buffer = export_mcqs_to_docx(mcqs)
        st.download_button("Download as DOCX", docx_buffer, file_name="generated_mcqs.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        html_content = export_mcqs_to_html(mcqs)
        st.download_button("Download as HTML", html_content, file_name="generated_mcqs.html", mime="text/html")